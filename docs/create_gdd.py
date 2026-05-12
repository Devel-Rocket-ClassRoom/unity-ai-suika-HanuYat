from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

doc = Document()

# ── 스타일 헬퍼 ──────────────────────────────────────────────────────────────
def set_heading(doc, text, level, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_para(doc, text, bold=False, italic=False, size=11, indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    return p

def add_table_row(table, cells, header=False):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = text
        if header:
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_bg(cell, '2E7D32')
    return row

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

# ── 제목 페이지 ──────────────────────────────────────────────────────────────
title = doc.add_heading('', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('수박 게임 (Suika Game)')
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(46, 125, 50)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Game Design Document v1.0')
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(100, 100, 100)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = meta.add_run(f'프로젝트: unity-ai-suika-HanuYat  |  작성일: {datetime.date.today().strftime("%Y년 %m월 %d일")}')
r2.font.size = Pt(11)
r2.font.color.rgb = RGBColor(120, 120, 120)
r2.italic = True

doc.add_page_break()

# ── 목차 ────────────────────────────────────────────────────────────────────
set_heading(doc, '목차', 1, (46, 125, 50))
toc_items = [
    '1. 게임 개요',
    '2. 핵심 게임플레이',
    '3. 과일 진화 시스템',
    '4. 점수 시스템',
    '5. UI/UX 설계',
    '6. Unity AI 활용 계획',
    '7. 기술 아키텍처',
    '8. 씬 구성',
    '9. 개발 일정 (당일 완성 계획)',
    '10. 리스크 및 대응',
]
for item in toc_items:
    add_para(doc, item, size=12)

doc.add_page_break()

# ── 1. 게임 개요 ─────────────────────────────────────────────────────────────
set_heading(doc, '1. 게임 개요', 1, (46, 125, 50))

set_heading(doc, '1.1 기본 정보', 2)
info_table = doc.add_table(rows=1, cols=2)
info_table.style = 'Table Grid'
headers = ['항목', '내용']
for i, h in enumerate(headers):
    cell = info_table.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.bold = True
    set_cell_bg(cell, '2E7D32')
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(255, 255, 255)

rows = [
    ('게임 제목', '수박 게임 (Suika Game) - HanuYat Edition'),
    ('장르', '퍼즐 / 캐주얼 / 물리 기반 병합'),
    ('플랫폼', 'PC (Windows), 향후 Mobile 확장 가능'),
    ('엔진', 'Unity 2022+ / URP 2D'),
    ('AI 도구', 'Unity AI (Muse Sprite·Texture·Animation Generator, AI Assistant)'),
    ('개발 기간', '1일 (오늘 완성 목표)'),
    ('대상 연령', '전연령'),
    ('조작', '마우스 클릭 / 터치'),
]
for label, value in rows:
    row = info_table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value
    for run in row.cells[0].paragraphs[0].runs:
        run.bold = True

doc.add_paragraph()

set_heading(doc, '1.2 게임 콘셉트', 2)
add_para(doc,
    '수박 게임은 2021년 일본 개발사 Aladdin X가 출시하고 닌텐도 스위치를 통해 전 세계적으로 바이럴된 '
    '과일 병합 퍼즐 게임의 Unity 재현작입니다. '
    '플레이어는 상자 안으로 과일을 떨어뜨려 동일한 과일끼리 충돌·병합시키며 점수를 쌓습니다. '
    '단순하지만 중독성 높은 물리 기반 퍼즐로, 11단계의 과일 진화를 통해 수박(최고 등급)을 만드는 것이 목표입니다.')

set_heading(doc, '1.3 핵심 재미 요소', 2)
add_bullet(doc, '예측 불가한 물리 연쇄반응 — 과일이 굴러다니며 예상치 못한 병합 발생')
add_bullet(doc, '단계적 성장감 — 체리→수박까지 11단계 진화의 성취감')
add_bullet(doc, '원 모어 턴 중독성 — 실패 직전 아슬아슬한 긴장감')
add_bullet(doc, 'AI 생성 아트 — Unity AI로 생성한 유니크한 과일 스프라이트')

doc.add_page_break()

# ── 2. 핵심 게임플레이 ────────────────────────────────────────────────────────
set_heading(doc, '2. 핵심 게임플레이', 1, (46, 125, 50))

set_heading(doc, '2.1 기본 규칙', 2)
add_para(doc, '① 플레이어는 상자 상단에서 과일을 좌우로 이동시켜 원하는 위치에 떨어뜨립니다.')
add_para(doc, '② 동일한 과일 두 개가 충돌하면 병합되어 다음 단계의 더 큰 과일이 됩니다.')
add_para(doc, '③ 병합 시 점수가 추가됩니다. (더 큰 과일일수록 높은 점수)')
add_para(doc, '④ 과일이 상자 위쪽 경계선을 넘으면 게임 오버.')
add_para(doc, '⑤ 다음에 떨어질 과일을 미리 1개 예고로 표시합니다.')
add_para(doc, '⑥ 플레이어가 투하할 수 있는 과일은 1~5단계(체리~감)로 제한됩니다.')

set_heading(doc, '2.2 물리 메커니즘', 2)
add_bullet(doc, 'Unity Physics 2D (Rigidbody2D + CircleCollider2D) 기반 물리 시뮬레이션')
add_bullet(doc, '과일은 완전한 원형으로 표현, 물리적으로 굴러다님')
add_bullet(doc, '병합 발생 지점: 두 과일의 충돌 중심점')
add_bullet(doc, '병합 시 이전 두 과일 제거 후 새 과일 Instantiate (약간의 파티클 이펙트)')
add_bullet(doc, '수박(11단계)끼리 병합 시 두 수박 제거 + 대규모 점수 보너스')

set_heading(doc, '2.3 조작 방법', 2)
ctrl_table = doc.add_table(rows=1, cols=3)
ctrl_table.style = 'Table Grid'
for i, h in enumerate(['조작', '입력', '설명']):
    cell = ctrl_table.rows[0].cells[i]
    cell.text = h
    set_cell_bg(cell, '2E7D32')
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

for row_data in [
    ('과일 이동', '마우스 X 이동', '떨어뜨릴 위치 결정'),
    ('과일 투하', '마우스 좌클릭', '현재 위치에서 과일 낙하'),
    ('다음 과일 확인', 'UI 미리보기', '우측 상단 Next 패널'),
]:
    r = ctrl_table.add_row()
    for i, text in enumerate(row_data):
        r.cells[i].text = text

doc.add_paragraph()

set_heading(doc, '2.4 게임 오버 조건', 2)
add_bullet(doc, '과일이 상자의 상단 위험선(빨간 점선)을 벗어난 상태로 3초 이상 유지될 때 게임 오버')
add_bullet(doc, '게임 오버 화면: 최종 점수, 최고 기록 갱신 여부, 재시작 버튼 표시')

doc.add_page_break()

# ── 3. 과일 진화 시스템 ───────────────────────────────────────────────────────
set_heading(doc, '3. 과일 진화 시스템', 1, (46, 125, 50))

set_heading(doc, '3.1 11단계 진화 체인', 2)
fruit_table = doc.add_table(rows=1, cols=5)
fruit_table.style = 'Table Grid'
for i, h in enumerate(['단계', '과일명 (KR)', '과일명 (EN)', '반지름 (px)', '병합 점수']):
    cell = fruit_table.rows[0].cells[i]
    cell.text = h
    set_cell_bg(cell, '2E7D32')
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

fruits = [
    ('1', '체리',     'Cherry',      '~24',  '1'),
    ('2', '딸기',     'Strawberry',  '~32',  '3'),
    ('3', '포도',     'Grape',       '~40',  '6'),
    ('4', '데코폰',   'Dekopon',     '~50',  '10'),
    ('5', '감',       'Persimmon',   '~60',  '15'),
    ('6', '사과',     'Apple',       '~72',  '21'),
    ('7', '배',       'Pear',        '~88',  '28'),
    ('8', '복숭아',   'Peach',       '~100', '36'),
    ('9', '파인애플', 'Pineapple',   '~116', '45'),
    ('10','멜론',     'Melon',       '~132', '55'),
    ('11','수박',     'Watermelon',  '~152', '66 + 보너스'),
]
for row_data in fruits:
    r = fruit_table.add_row()
    for i, text in enumerate(row_data):
        r.cells[i].text = text

doc.add_paragraph()

set_heading(doc, '3.2 투하 가능 과일', 2)
add_para(doc, '플레이어가 직접 투하할 수 있는 과일은 1단계(체리) ~ 5단계(감)으로 제한됩니다.')
add_para(doc, '각 턴마다 랜덤으로 결정되며, 다음 차례 과일은 UI에 미리 표시됩니다.')

set_heading(doc, '3.3 Unity AI 스프라이트 생성', 2)
add_para(doc, 'Unity AI Generators (Sprite Generator)를 활용하여 각 과일의 2D 스프라이트를 AI로 생성합니다.')
add_bullet(doc, '프롬프트 예시: "cute cartoon cherry fruit, 2D game sprite, white background, vibrant colors"')
add_bullet(doc, 'LoRA 모델: Scenario / Layer (Stable Diffusion / Flux 기반)')
add_bullet(doc, '생성 후 Unity Sprite Editor로 트리밍 및 Pivot 조정')
add_bullet(doc, '필요시 Unity AI Texture Generator로 배경·컨테이너 텍스처도 생성')

doc.add_page_break()

# ── 4. 점수 시스템 ────────────────────────────────────────────────────────────
set_heading(doc, '4. 점수 시스템', 1, (46, 125, 50))

set_heading(doc, '4.1 기본 점수', 2)
add_para(doc, '병합 발생 시: 해당 단계 과일의 병합 점수 즉시 추가 (위 표 참조)')

set_heading(doc, '4.2 보너스 점수', 2)
add_bullet(doc, '콤보 보너스: 연속 병합 발생 시 ×1.5 → ×2.0 → ×3.0 배율 적용')
add_bullet(doc, '수박 달성 보너스: 수박 생성 시 +200점 추가')
add_bullet(doc, '수박 병합 보너스: 수박 두 개 병합 성공 시 +500점 추가')

set_heading(doc, '4.3 점수 저장', 2)
add_bullet(doc, '현재 세션 최고점: PlayerPrefs로 로컬 저장')
add_bullet(doc, '점수 표시: 상단 HUD에 실시간 갱신')

doc.add_page_break()

# ── 5. UI/UX 설계 ─────────────────────────────────────────────────────────────
set_heading(doc, '5. UI/UX 설계', 1, (46, 125, 50))

set_heading(doc, '5.1 화면 레이아웃', 2)
layout_text = """
┌─────────────────────────────────────┐
│  SCORE: 0000   │  BEST: 0000        │
│────────────────│────────────────────│
│                │  [NEXT]            │
│                │  ┌──────┐          │
│                │  │  🍒  │          │
│  ┌───────────┐ │  └──────┘          │
│  │           │ │                    │
│  │  게임 영역  │ │                    │
│  │  (상자)   │ │                    │
│  │           │ │                    │
│  │  과일들   │ │                    │
│  └───────────┘ │                    │
└─────────────────────────────────────┘
"""
p = doc.add_paragraph()
run = p.add_run(layout_text)
run.font.name = 'Courier New'
run.font.size = Pt(9)

set_heading(doc, '5.2 HUD 요소', 2)
add_bullet(doc, '현재 점수 (좌측 상단)')
add_bullet(doc, '최고 점수 (우측 상단)')
add_bullet(doc, 'Next 과일 미리보기 (우측 패널)')
add_bullet(doc, '위험선 표시 (상자 상단 빨간 점선)')
add_bullet(doc, '게임 오버 오버레이 (점수 + 재시작 버튼)')

set_heading(doc, '5.3 씬 전환', 2)
add_bullet(doc, '메인 메뉴 → 게임 씬: 페이드 인/아웃')
add_bullet(doc, '게임 오버 → 재시작: DOTween 또는 Unity 기본 애니메이션')

doc.add_page_break()

# ── 6. Unity AI 활용 계획 ────────────────────────────────────────────────────
set_heading(doc, '6. Unity AI 활용 계획', 1, (46, 125, 50))

set_heading(doc, '6.1 Unity AI 도구 목록', 2)
ai_table = doc.add_table(rows=1, cols=3)
ai_table.style = 'Table Grid'
for i, h in enumerate(['Unity AI 기능', '용도', '적용 대상']):
    cell = ai_table.rows[0].cells[i]
    cell.text = h
    set_cell_bg(cell, '2E7D32')
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

ai_rows = [
    ('Sprite Generator', '과일 스프라이트 생성', '체리~수박 11종'),
    ('Texture Generator', '배경·상자 텍스처', '게임 배경, 나무 상자'),
    ('Animation Generator', '과일 등장·병합 애니', '스폰·병합·팝 이펙트'),
    ('AI Assistant (Chat)', '코드 생성 / 버그 수정', 'C# 스크립트 전반'),
    ('Claude Code + Unity MCP', '씬/오브젝트 자동 구성', 'GameObjects, Scripts'),
]
for row_data in ai_rows:
    r = ai_table.add_row()
    for i, text in enumerate(row_data):
        r.cells[i].text = text

doc.add_paragraph()

set_heading(doc, '6.2 Claude Code + Unity MCP 워크플로우', 2)
add_para(doc, 'Claude Code는 Unity MCP 서버를 통해 Unity Editor와 직접 통신합니다.')
steps = [
    '1. Unity MCP로 씬 오브젝트 생성 및 컴포넌트 자동 배치',
    '2. C# 스크립트 자동 생성 (FruitSpawner, FruitMerger, ScoreManager 등)',
    '3. 생성된 스프라이트를 자동으로 프리팹에 할당',
    '4. 게임 로직 검증 및 콘솔 로그 모니터링',
    '5. 씬 뷰 캡처로 레이아웃 확인',
]
for step in steps:
    add_para(doc, step, indent=0.5)

set_heading(doc, '6.3 AI 생성 프롬프트 예시', 2)
prompts = [
    ('체리 스프라이트', 'cute cartoon cherry, 2D game sprite, isolated on white, vibrant red, glossy, simple style'),
    ('딸기 스프라이트', 'cute kawaii strawberry, 2D game art, isolated white background, bright colors, round shape'),
    ('수박 스프라이트', 'cute round watermelon, 2D game sprite, white background, cartoon style, green and red'),
    ('게임 배경 텍스처', 'wooden game board texture, top-down 2D game, warm brown tones, seamless'),
]
p_table = doc.add_table(rows=1, cols=2)
p_table.style = 'Table Grid'
for i, h in enumerate(['과일', '생성 프롬프트']):
    cell = p_table.rows[0].cells[i]
    cell.text = h
    set_cell_bg(cell, '2E7D32')
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
for fruit, prompt in prompts:
    r = p_table.add_row()
    r.cells[0].text = fruit
    r.cells[1].text = prompt

doc.add_paragraph()
doc.add_page_break()

# ── 7. 기술 아키텍처 ─────────────────────────────────────────────────────────
set_heading(doc, '7. 기술 아키텍처', 1, (46, 125, 50))

set_heading(doc, '7.1 스크립트 목록', 2)
script_table = doc.add_table(rows=1, cols=3)
script_table.style = 'Table Grid'
for i, h in enumerate(['스크립트명', '역할', '주요 메서드']):
    cell = script_table.rows[0].cells[i]
    cell.text = h
    set_cell_bg(cell, '2E7D32')
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

scripts = [
    ('GameManager', '게임 상태 관리 (시작/종료/재시작)', 'StartGame(), GameOver(), RestartGame()'),
    ('FruitSpawner', '과일 생성 및 투하 제어', 'SpawnFruit(), DropFruit(), MovePreview()'),
    ('Fruit', '개별 과일 데이터·충돌 감지', 'OnCollisionEnter2D(), GetFruitData()'),
    ('FruitMerger', '병합 로직 처리', 'TryMerge(), CreateMergedFruit()'),
    ('ScoreManager', '점수 계산 및 저장', 'AddScore(), SaveHighScore()'),
    ('UIManager', 'HUD 업데이트', 'UpdateScore(), ShowGameOver()'),
    ('FruitData', '과일 ScriptableObject 정의', '(데이터 컨테이너)'),
]
for row_data in scripts:
    r = script_table.add_row()
    for i, text in enumerate(row_data):
        r.cells[i].text = text

doc.add_paragraph()

set_heading(doc, '7.2 프리팹 구조', 2)
add_para(doc, 'Fruit Prefab 구성:')
add_bullet(doc, 'Rigidbody2D: gravity scale=1.0, linear drag=0.5')
add_bullet(doc, 'CircleCollider2D: radius = 과일 크기에 맞춤')
add_bullet(doc, 'SpriteRenderer: Unity AI 생성 스프라이트 할당')
add_bullet(doc, 'Fruit.cs 스크립트: fruitLevel, score, nextFruitPrefab')

set_heading(doc, '7.3 씬 오브젝트 계층', 2)
hierarchy = """
[Scene: Suika]
├── GameManager (GameManager.cs)
├── Camera (Main Camera)
├── Canvas (UI)
│   ├── ScoreText
│   ├── BestScoreText
│   ├── NextFruitPanel
│   └── GameOverPanel
├── GameArea
│   ├── WallLeft
│   ├── WallRight
│   ├── WallBottom
│   └── DangerLine
└── FruitSpawner (FruitSpawner.cs)
"""
p = doc.add_paragraph()
run = p.add_run(hierarchy)
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_page_break()

# ── 8. 씬 구성 ───────────────────────────────────────────────────────────────
set_heading(doc, '8. 씬 구성', 1, (46, 125, 50))

set_heading(doc, '8.1 씬 목록', 2)
add_bullet(doc, 'Assets/Scenes/Suika.unity — 메인 게임 씬 (기존 생성됨)')
add_bullet(doc, 'Assets/Scenes/MainMenu.unity — 메인 메뉴 씬 (필요시)')

set_heading(doc, '8.2 카메라 설정', 2)
add_bullet(doc, '2D Orthographic Camera, Size = 6')
add_bullet(doc, '해상도: 9:16 세로 기준 (540×960)')
add_bullet(doc, 'URP 2D Renderer 사용 (기존 프로젝트 설정 유지)')

set_heading(doc, '8.3 물리 설정', 2)
add_bullet(doc, 'Physics 2D Gravity: (0, -9.81)')
add_bullet(doc, 'Layer: Fruit, Wall — 충돌 레이어 분리')
add_bullet(doc, 'PhysicsMaterial2D: bounciness=0.3, friction=0.5')

doc.add_page_break()

# ── 9. 개발 일정 ─────────────────────────────────────────────────────────────
set_heading(doc, '9. 개발 일정 (당일 완성 계획)', 1, (46, 125, 50))

schedule_table = doc.add_table(rows=1, cols=3)
schedule_table.style = 'Table Grid'
for i, h in enumerate(['시간대', '작업 항목', '담당 / 도구']):
    cell = schedule_table.rows[0].cells[i]
    cell.text = h
    set_cell_bg(cell, '2E7D32')
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

schedule = [
    ('Phase 1\n(GDD·설계)', 'GDD 작성, 씬 설계, 프리팹 구조 확정', 'Claude Code (완료)'),
    ('Phase 2\n(AI 아트)', 'Unity AI Sprite Generator로 11종 과일 생성\n배경·상자 텍스처 생성', 'Unity AI Generators'),
    ('Phase 3\n(기본 로직)', 'FruitData ScriptableObject 생성\nFruitSpawner, Fruit, FruitMerger 스크립트 작성', 'Claude Code + Unity MCP'),
    ('Phase 4\n(UI)', 'Canvas HUD 구성\nScore, Next, GameOver 패널 완성', 'Unity MCP + UI Builder'),
    ('Phase 5\n(물리·병합)', '물리 튜닝, 병합 로직 검증\n점수 시스템 연결', 'Unity Editor + Claude Code'),
    ('Phase 6\n(이펙트·사운드)', '병합 파티클, UI 애니, 기본 SFX 추가', 'Unity Particle System'),
    ('Phase 7\n(테스트·빌드)', '플레이 테스트, 버그 수정, PC 빌드', 'Unity Build'),
]
for row_data in schedule:
    r = schedule_table.add_row()
    for i, text in enumerate(row_data):
        r.cells[i].text = text

doc.add_paragraph()
doc.add_page_break()

# ── 10. 리스크 및 대응 ───────────────────────────────────────────────────────
set_heading(doc, '10. 리스크 및 대응', 1, (46, 125, 50))

risk_table = doc.add_table(rows=1, cols=3)
risk_table.style = 'Table Grid'
for i, h in enumerate(['리스크', '영향도', '대응 방안']):
    cell = risk_table.rows[0].cells[i]
    cell.text = h
    set_cell_bg(cell, '2E7D32')
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

risks = [
    ('Unity AI 크레딧 부족', '중', 'Free Trial 1,000 크레딧 활용; 부족 시 무료 오픈소스 과일 아이콘 대체'),
    ('물리 병합 버그 (무한 병합 루프)', '높음', '병합 중 플래그(isMerging) 설정으로 중복 충돌 방지'),
    ('성능 저하 (과일 과다)', '낮음', '상자 크기 제한으로 최대 오브젝트 수 자연 제한'),
    ('당일 완성 시간 부족', '중', 'SFX·애니메이션은 선택 사항으로 분류; 핵심 루프 우선'),
    ('Unity MCP 연결 문제', '낮음', 'Unity Editor 재시작 후 MCP 재연결; 수동 스크립트 작업으로 대체'),
]
for row_data in risks:
    r = risk_table.add_row()
    for i, text in enumerate(row_data):
        r.cells[i].text = text

doc.add_paragraph()

# ── 참고 자료 ─────────────────────────────────────────────────────────────────
set_heading(doc, '참고 자료', 1, (46, 125, 50))
refs = [
    'Suika Game - Wikipedia: https://en.wikipedia.org/wiki/Suika_Game',
    'Suika Game Fandom Wiki (Fruit List): https://suikagame.fandom.com/wiki/List_of_fruits',
    'Unity AI Official: https://unity.com/features/ai',
    'Unity Muse: https://muse.unity.com/',
    'Unity AI in Unity 6.2 - CG Channel: https://www.cgchannel.com/2025/08/unity-rolls-out-unity-ai-in-unity-6-2/',
]
for ref in refs:
    add_bullet(doc, ref)

# ── 저장 ─────────────────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'GDD_수박게임_HanuYat.docx')
doc.save(output_path)
print(f'GDD 저장 완료: {output_path}')
